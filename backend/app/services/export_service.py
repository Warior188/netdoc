import os
from datetime import datetime
from typing import List
from pathlib import Path

from app.schemas.schemas import ProjectDetail, DeviceConfig
from app.core.config import settings


def _build_html(project: ProjectDetail, configs: List[DeviceConfig]) -> str:
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    addr_rows = ""
    for entry in project.address_entries:
        addr_rows += f"""
        <tr>
            <td>{entry.device_name}</td>
            <td>{entry.interface}</td>
            <td>{entry.ip_address or '—'}</td>
            <td>{entry.subnet_mask}</td>
            <td>{entry.gateway or '—'}</td>
            <td>{entry.vlan or '—'}</td>
        </tr>"""

    devices_rows = ""
    for d in project.devices:
        devices_rows += f"""
        <tr>
            <td>{d.name}</td>
            <td>{d.device_type.value}</td>
        </tr>"""

    config_sections = ""
    for cfg in configs:
        config_sections += f"""
        <div class="config-section">
            <h3>{cfg.device_name} <span class="type-badge">{cfg.device_type.value}</span></h3>
            <pre><code>{cfg.config_text}</code></pre>
        </div>"""

    return f"""<!DOCTYPE html>
<html lang="pl">
<head>
<meta charset="UTF-8">
<style>
  @import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;700&family=IBM+Plex+Sans:wght@400;500;600&display=swap');
  body {{ font-family: 'IBM Plex Sans', sans-serif; color: #1a1a1a; margin: 0; padding: 40px; font-size: 12px; }}
  h1 {{ font-size: 24px; font-weight: 600; border-bottom: 2px solid #042C53; padding-bottom: 8px; }}
  h2 {{ font-size: 16px; font-weight: 600; color: #042C53; margin-top: 32px; border-left: 4px solid #378ADD; padding-left: 10px; }}
  h3 {{ font-size: 13px; font-weight: 500; margin-bottom: 6px; }}
  .meta {{ color: #666; font-size: 11px; margin-top: 4px; }}
  table {{ width: 100%; border-collapse: collapse; margin: 12px 0; font-size: 11px; }}
  th {{ background: #042C53; color: white; padding: 7px 10px; text-align: left; }}
  td {{ padding: 6px 10px; border-bottom: 1px solid #e0e0e0; }}
  tr:nth-child(even) td {{ background: #f7f8fa; }}
  pre {{ background: #1e2124; color: #97C459; padding: 16px; border-radius: 6px; font-family: 'JetBrains Mono', monospace; font-size: 10px; line-height: 1.6; overflow-x: auto; white-space: pre-wrap; word-break: break-all; }}
  .config-section {{ margin-bottom: 24px; page-break-inside: avoid; }}
  .type-badge {{ background: #E6F1FB; color: #0C447C; font-size: 10px; padding: 2px 8px; border-radius: 4px; font-family: 'JetBrains Mono', monospace; }}
  .page-break {{ page-break-before: always; }}
  footer {{ margin-top: 40px; color: #999; font-size: 10px; border-top: 1px solid #eee; padding-top: 12px; }}
</style>
</head>
<body>
<h1>Dokumentacja sieci: {project.name}</h1>
<p class="meta">Wygenerowano: {now} | NetDoc Maker v1.0</p>
{f'<p>{project.description}</p>' if project.description else ''}

<h2>1. Urządzenia sieciowe</h2>
<table>
  <thead><tr><th>Nazwa urządzenia</th><th>Typ</th></tr></thead>
  <tbody>{devices_rows}</tbody>
</table>

<h2>2. Tablica adresacji IP</h2>
<table>
  <thead>
    <tr><th>Urządzenie</th><th>Interfejs</th><th>Adres IP</th><th>Maska podsieci</th><th>Brama domyślna</th><th>VLAN</th></tr>
  </thead>
  <tbody>{addr_rows or '<tr><td colspan="6">Brak wpisów adresacji</td></tr>'}</tbody>
</table>

<div class="page-break"></div>

<h2>3. Konfiguracje urządzeń</h2>
{config_sections or '<p>Brak konfiguracji do wygenerowania.</p>'}

<footer>NetDoc Maker | Dokumentacja wygenerowana automatycznie | {now}</footer>
</body>
</html>"""


def export_pdf(project: ProjectDetail, configs: List[DeviceConfig]) -> str:
    from weasyprint import HTML

    html_content = _build_html(project, configs)
    filename = f"netdoc_{project.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    filepath = os.path.join(settings.EXPORT_DIR, filename)

    HTML(string=html_content).write_pdf(filepath)
    return filename


def export_docx(project: ProjectDetail, configs: List[DeviceConfig]) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading(f"Dokumentacja sieci: {project.name}", 0)
    title.runs[0].font.color.rgb = RGBColor(0x04, 0x2C, 0x53)

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    doc.add_paragraph(f"Wygenerowano: {now} | NetDoc Maker v1.0")

    if project.description:
        doc.add_paragraph(project.description)

    doc.add_heading("1. Urządzenia sieciowe", level=1)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Nazwa urządzenia"
    hdr[1].text = "Typ"
    for d in project.devices:
        row = tbl.add_row().cells
        row[0].text = d.name
        row[1].text = d.device_type.value

    doc.add_heading("2. Tablica adresacji IP", level=1)
    cols = ["Urządzenie", "Interfejs", "Adres IP", "Maska", "Brama", "VLAN"]
    tbl2 = doc.add_table(rows=1, cols=len(cols))
    tbl2.style = "Table Grid"
    hdr2 = tbl2.rows[0].cells
    for i, c in enumerate(cols):
        hdr2[i].text = c
    for entry in project.address_entries:
        row = tbl2.add_row().cells
        row[0].text = entry.device_name
        row[1].text = entry.interface
        row[2].text = entry.ip_address or "—"
        row[3].text = entry.subnet_mask
        row[4].text = entry.gateway or "—"
        row[5].text = str(entry.vlan) if entry.vlan else "—"

    doc.add_heading("3. Konfiguracje urządzeń", level=1)
    for cfg in configs:
        doc.add_heading(f"{cfg.device_name} ({cfg.device_type.value})", level=2)
        p = doc.add_paragraph()
        run = p.add_run(cfg.config_text)
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        p.paragraph_format.space_after = Pt(12)

    filename = f"netdoc_{project.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(settings.EXPORT_DIR, filename)
    doc.save(filepath)
    return filename